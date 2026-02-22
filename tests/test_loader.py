"""Tests for document loader."""

import tempfile
from pathlib import Path
from src.ingestion.loader import load_txt, load_markdown, load_docx, load_documents


def test_load_txt():
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
        f.write("Hello, this is a test document.")
        f.flush()
        doc = load_txt(Path(f.name))
        assert doc.content == "Hello, this is a test document."
        assert doc.format == "txt"


def test_load_markdown():
    with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False) as f:
        f.write("# Title\n\nSome **bold** text.")
        f.flush()
        doc = load_markdown(Path(f.name))
        assert "Title" in doc.content
        assert "bold" in doc.content
        assert doc.format == "md"


def test_load_documents_from_directory():
    with tempfile.TemporaryDirectory() as tmpdir:
        # Create a test file
        p = Path(tmpdir) / "test.txt"
        p.write_text("Test content")

        docs = load_documents(Path(tmpdir))
        assert len(docs) == 1
        assert docs[0].content == "Test content"


def test_load_docx():
    """Test loading a .docx file."""
    from docx import Document as DocxDocument

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        tmp_path = Path(f.name)
    # Create a minimal .docx using python-docx
    doc = DocxDocument()
    doc.add_paragraph("Hello from a Word document.")
    doc.add_paragraph("This is the second paragraph.")
    doc.save(str(tmp_path))

    loaded = load_docx(tmp_path)
    assert "Hello from a Word document." in loaded.content
    assert "second paragraph" in loaded.content
    assert loaded.format == "docx"
    tmp_path.unlink()
