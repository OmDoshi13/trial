"""Multi-format document loader (PDF, TXT, Markdown, DOC/DOCX).

Each format has its own loader function — add a new one to support more types.
"""

from pathlib import Path
from dataclasses import dataclass
import fitz  # PyMuPDF
import markdown
from bs4 import BeautifulSoup
from docx import Document as DocxDocument


@dataclass
class Document:
    content: str
    source: str
    format: str
    title: str


def load_pdf(path: Path) -> Document:
    """Extract text from PDF using PyMuPDF."""
    doc = fitz.open(str(path))
    text_parts = [page.get_text() for page in doc]
    doc.close()
    return Document(
        content="\n".join(text_parts).strip(),
        source=str(path), format="pdf", title=path.stem,
    )


def load_txt(path: Path) -> Document:
    content = path.read_text(encoding="utf-8")
    return Document(content=content.strip(), source=str(path), format="txt", title=path.stem)


def load_markdown(path: Path) -> Document:
    """Convert Markdown → HTML → plain text."""
    raw = path.read_text(encoding="utf-8")
    html = markdown.markdown(raw, extensions=["tables", "fenced_code"])
    text = BeautifulSoup(html, "html.parser").get_text(separator="\n")
    return Document(content=text.strip(), source=str(path), format="md", title=path.stem)


def load_docx(path: Path) -> Document:
    """Read paragraphs from a .docx file."""
    doc = DocxDocument(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return Document(
        content="\n".join(paragraphs).strip(),
        source=str(path), format="docx", title=path.stem,
    )


def load_doc(path: Path) -> Document:
    """Try reading .doc as docx first; fall back to raw text extraction."""
    try:
        return load_docx(path)
    except Exception:
        raw = path.read_bytes()
        text = raw.decode("utf-8", errors="ignore")
        lines = []
        for line in text.split("\n"):
            printable = sum(1 for c in line if c.isprintable() or c in ("\t",))
            if len(line) > 0 and printable / len(line) > 0.8:
                lines.append(line.strip())
        return Document(
            content="\n".join(lines).strip(),
            source=str(path), format="doc", title=path.stem,
        )


# Extension → loader mapping
LOADERS = {
    ".pdf": load_pdf,
    ".txt": load_txt,
    ".md": load_markdown,
    ".docx": load_docx,
    ".doc": load_doc,
}

SUPPORTED_EXTENSIONS = set(LOADERS.keys())


def load_single_file(path: Path) -> Document:
    """Load one file. Raises ValueError for unsupported/empty files."""
    ext = path.suffix.lower()
    loader = LOADERS.get(ext)
    if not loader:
        raise ValueError(
            f"Unsupported file format: '{ext}'. "
            f"Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
        )
    doc = loader(path)
    if not doc.content:
        raise ValueError(f"Document is empty: {path.name}")
    return doc


def load_documents(directory: Path) -> list[Document]:
    """Load all supported files from a directory, skip the rest."""
    documents = []
    if not directory.exists():
        raise FileNotFoundError(f"Documents directory not found: {directory}")

    for file_path in sorted(directory.iterdir()):
        ext = file_path.suffix.lower()
        loader = LOADERS.get(ext)
        if loader:
            print(f"  Loading [{ext}] {file_path.name}")
            try:
                doc = loader(file_path)
                if doc.content:
                    documents.append(doc)
                else:
                    print(f"  ⚠ Skipping empty: {file_path.name}")
            except Exception as e:
                print(f"  ✗ Error loading {file_path.name}: {e}")
        else:
            print(f"  ⊘ Skipping unsupported: {file_path.name}")

    return documents
